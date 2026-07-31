# -*- coding: utf-8 -*-
"""Build the bilingual (Hebrew RTL + English) AgroData strategy document
   reconstructed from the 3 Ofakim-2035 workshop photos in this folder."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
GREEN = RGBColor(0x2E, 0x6B, 0x3E)
DARK  = RGBColor(0x22, 0x2A, 0x24)

doc = Document()
# base font
st = doc.styles['Normal']
st.font.name = 'Arial'
st.font.size = Pt(11)

def rtl_p(p):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:bidi'); pPr.append(b)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def rtl_run(r):
    rPr = r._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl'); rPr.append(rtl)

def heading(text, rtl=False, color=GREEN, size=15, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    if rtl: rtl_p(p); rtl_run(r)
    return p

def para(text, rtl=False, bullet=False, size=11, color=DARK, bold=False):
    p = doc.add_paragraph(style='List Bullet' if bullet else None)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color; r.bold = bold
    if rtl: rtl_p(p); rtl_run(r)
    return p

def title_block(main, sub):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(main); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = GREEN
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(sub); r2.font.size = Pt(13); r2.font.color.rgb = RGBColor(0x6b,0x6f,0x59)

# ---------- COVER ----------
title_block("AgroData", "פלטפורמת נתונים חקלאית לאומית  ·  National Agricultural Data Platform")
n = doc.add_paragraph(); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = n.add_run("מסמך אסטרטגיה — טיוטה ראשונה, שוחזרה מדפי הסדנה (אופקים 2035)  ·  30/07/2026")
rn.italic = True; rn.font.size = Pt(9.5); rn.font.color.rgb = RGBColor(0x8a,0x7a,0x60)

# ======================================================= HEBREW (RTL)
heading("חלק א׳ — עברית", rtl=True, color=DARK, size=13)

heading("חזון", rtl=True)
para("הקמת פלטפורמת נתונים חקלאית לאומית (AgroData) שתמצב את ישראל כמרכז מוביל לנתונים "
     "ולחדשנות בחקלאות — מאגר לאומי שמושך חברות, סטארטאפים, אקדמיה ומשקיעים.", rtl=True)

heading("1. התוצאה שנרצה לקדם")  # keep number LTR-friendly
para("יצירת ״מנוע משיכה״ לחברות בתחום ה-AgriTech באמצעות מאגר נתונים חקלאי לאומי (פתוח/מבוקר) "
     "שמאיץ פיתוח, שיתופי פעולה ומחקר, ומחזק את הריבונות והחוסן של המידע החקלאי.", rtl=True)

heading("2. מהלכי העבודה המרכזיים")
for t in ["מאגר נתונים לאומי (Data Lake) — איסוף, אחסון וניהול נתוני חקלאות",
          "חיישנים ו-IoT בשטח — איסוף נתונים בזמן אמת",
          "Hub לאומי לחקלאות — נקודת מפגש לחברות, מגדלים ומחקר",
          "אקסלרטור / זרוע הפקה — האצת סטארטאפים ופיילוטים",
          "שכבת אנליטיקת AI — תובנות, חיזוי וקבלת החלטות",
          "שיתופי פעולה — חברות, אקדמיה, משרדי ממשלה ורשויות"]:
    para(t, rtl=True, bullet=True)

heading("3. מסלול היישום (השלבים)")
for t in ["שלב 1 — POC לאיסוף נתונים (הוכחת היתכנות ראשונית)",
          "שלב 2 — POC מורחב ואפיון המאגר",
          "שלב 3 — פיילוט בקנה מידה",
          "שלב 4 — הרחבה לאומית",
          "שלב 5 — גיוס מימון והמשכיות"]:
    para(t, rtl=True, bullet=True)

heading("4. תשתיות תומכות")
for t in ["מדיניות ורגולציה (Policy) — ממשל נתונים, פרטיות ובעלות",
          "צוות מוביל (Team) — ליבה רב-תחומית",
          "מודל מימון — ציבורי/פרטי, מענקים ומשקיעים",
          "חוות דעת ואפיון — היתכנות טכנית וכלכלית"]:
    para(t, rtl=True, bullet=True)

heading("5. איך יוצאים לדרך?")
para("טבלת פעולות פותחת (למילוי מוביל/שותפים/יעד):", rtl=True)

# action table (bilingual headers)
rows = [
    ("פעולה / Action", "מוביל / Lead", "שותפים / Partners", "יעד / Target"),
    ("POC לאיסוף נתונים / Data-collection POC", "", "", ""),
    ("אפיון מאגר וארכיטקטורה / Repository & architecture", "", "", ""),
    ("פיילוט בקנה מידה / Scaled pilot", "", "", ""),
    ("גיבוש מדיניות ורגולציה / Policy & regulation", "", "", ""),
    ("הקמת אקסלרטור / Accelerator setup", "", "", ""),
]
tbl = doc.add_table(rows=len(rows), cols=4)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, htext in enumerate(rows[0]):
    c = tbl.rows[0].cells[ci]; c.paragraphs[0].add_run(htext).bold = True
for ri in range(1, len(rows)):
    for ci in range(4):
        tbl.rows[ri].cells[ci].text = rows[ri][ci]

doc.add_page_break()

# ======================================================= ENGLISH (LTR)
heading("Part B — English", color=DARK, size=13)

heading("Vision")
para("Establish a National Agricultural Data Platform (AgroData) that positions Israel as a "
     "leading hub for agricultural data and innovation — a national repository that attracts "
     "companies, startups, academia, and investors.")

heading("1. Desired Outcome")
para("Create an “attraction engine” for AgriTech companies through a national agricultural data "
     "repository (open / governed) that accelerates development, collaboration, and research, and "
     "strengthens the sovereignty and resilience of agricultural data.")

heading("2. Key Work Processes")
for t in ["National Data Lake — collect, store, and manage agricultural data",
          "Field sensors & IoT — real-time data capture",
          "National Agriculture Hub — meeting point for companies, growers, and research",
          "Accelerator / execution arm — accelerate startups and pilots",
          "AI analytics layer — insights, forecasting, and decision support",
          "Partnerships — companies, academia, government ministries and authorities"]:
    para(t, bullet=True)

heading("3. Implementation Roadmap")
for t in ["Stage 1 — Data-collection POC (initial proof of concept)",
          "Stage 2 — Expanded POC & repository characterization",
          "Stage 3 — Scaled pilot",
          "Stage 4 — National scale-up",
          "Stage 5 — Funding & sustainability"]:
    para(t, bullet=True)

heading("4. Enablers")
for t in ["Policy & regulation — data governance, privacy, and ownership",
          "Core team — multidisciplinary leadership",
          "Funding model — public/private, grants and investors",
          "Feasibility & characterization — technical and economic"]:
    para(t, bullet=True)

heading("5. How We Get Started")
para("See the opening action table in Part A (fill in Lead / Partners / Target).")

# note
doc.add_paragraph()
note = doc.add_paragraph()
rr = note.add_run("Note: first draft reconstructed from three handwritten workshop pages "
                  "(Ofakim 2035). Some handwriting was partly illegible — please review and complete.")
rr.italic = True; rr.font.size = Pt(9); rr.font.color.rgb = RGBColor(0x8a,0x7a,0x60)

out = os.path.join(HERE, "AgroData - אסטרטגיה HE-EN.docx")
doc.save(out)
print("SAVED:", out)
